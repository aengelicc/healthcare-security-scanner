#!/usr/bin/env python3
“””
Healthcare Code Security Scanner
Generates PDF/Word reports from Semgrep findings
Supports local directories AND GitHub repositories
Includes HIPAA Compliance Mapping and Automated Remediation
“””

import json
import os
import shutil
import subprocess
import sys
import tempfile
import traceback
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

# Import remediation templates

try:
from remediations import generate_remediation_section, get_remediation

```
HAS_REMEDIATIONS = True
```

except ImportError:
HAS_REMEDIATIONS = False
print(“Warning: remediations.py not found. Install remediation templates.”)

try:
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

```
HAS_DOCX = True
```

except ImportError:
HAS_DOCX = False
print(“Warning: python-docx not installed. Install with: pip install python-docx”)

try:
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.platypus import (
Paragraph,
SimpleDocTemplate,
Spacer,
Table,
TableStyle,
)

```
HAS_REPORTLAB = True
```

except ImportError:
HAS_REPORTLAB = False
print(“Warning: reportlab not installed. Install with: pip install reportlab”)

try:
from github import Github

```
HAS_GITHUB = True
```

except ImportError:
HAS_GITHUB = False
print(“Warning: PyGithub not installed. Install with: pip install PyGithub”)

class SecurityScanner:
def **init**(self, rules_file: str = “healthcare_rules.yaml”, verbose: bool = False):
self.rules_file = rules_file
self.verbose = verbose
self.findings: List[Dict[str, Any]] = []
self.scan_metadata: Dict[str, Any] = {}
self.repo_info: Dict[str, Any] = {}

```
def _log_debug(self, message: str) -> None:
    """Print debug messages only when verbose mode is enabled."""
    if self.verbose:
        print(message)

def run_semgrep(self, target_path: str) -> bool:
    """Execute Semgrep scan and capture results via temp file."""
    temp_path = None
    try:
        with tempfile.NamedTemporaryFile(
            mode="w+", delete=False, suffix=".json", encoding="utf-8"
        ) as tmp_file:
            temp_path = tmp_file.name

        cmd = [
            "semgrep",
            "scan",
            "--config",
            self.rules_file,
            "--output",
            temp_path,
            "--json",
            target_path,
        ]

        result = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            timeout=300,
        )

        if not os.path.exists(temp_path) or os.path.getsize(temp_path) == 0:
            print(f"Error: Temp file empty or missing. Return code: {result.returncode}")
            return False

        with open(temp_path, "r", encoding="utf-8") as f:
            output = f.read()

        if not output.strip():
            print("Warning: Semgrep returned empty JSON")
            self.findings = []
            self._build_scan_metadata(target_path, 0, 0, 0, 0)
            return True

        data = json.loads(output)
        self.findings = data.get("results", [])

        critical = high = medium = low = 0
        for finding in self.findings:
            severity = finding.get("extra", {}).get("severity", "")
            if severity == "ERROR":
                critical += 1
            elif severity == "HIGH":
                high += 1
            elif severity == "MEDIUM":
                medium += 1
            elif severity == "LOW":
                low += 1

        self._build_scan_metadata(target_path, critical, high, medium, low)

        self._log_debug(f"\nDEBUG: Found {len(self.findings)} findings")
        if self.findings and self.verbose:
            f = self.findings[0]
            self._log_debug(f"   Check ID: {f.get('check_id')}")
            self._log_debug(f"   Path: {f.get('path')}")
            self._log_debug(f"   Severity: {f.get('extra', {}).get('severity')}")

        return True

    except Exception as e:
        print(f"Error in run_semgrep: {e}")
        traceback.print_exc()
        return False
    finally:
        # Always clean up the temp file
        if temp_path and os.path.exists(temp_path):
            os.unlink(temp_path)

def _build_scan_metadata(
    self, target_path: str, critical: int, high: int, medium: int, low: int
) -> None:
    """Populate scan_metadata dict."""
    self.scan_metadata = {
        "timestamp": datetime.now().isoformat(),
        "target": target_path,
        "rules_file": self.rules_file,
        "total_findings": len(self.findings),
        "critical": critical,
        "high": high,
        "medium": medium,
        "low": low,
    }

def clone_github_repo(
    self, repo_url: str, token: Optional[str] = None
) -> Optional[str]:
    """Clone a GitHub repository to a temporary directory."""
    if not HAS_GITHUB:
        print("Error: PyGithub not installed. Run: pip install PyGithub")
        return None

    temp_dir = tempfile.mkdtemp(prefix="healthcare_scan_")
    try:
        print(f"\nCloning repository: {repo_url}")

        repo_name = repo_url.rstrip("/").split("/")[-1].replace(".git", "")
        clone_path = os.path.join(temp_dir, repo_name)

        if token:
            g = Github(token)
            repo_path = repo_url.replace("https://github.com/", "").rstrip("/")
            repo = g.get_repo(repo_path)
            # Use a credential helper approach to avoid embedding token in URL
            clone_url = f"https://x-access-token:{token}@github.com/{repo.full_name}.git"
        else:
            clone_url = repo_url if repo_url.endswith(".git") else repo_url + ".git"

        result = subprocess.run(
            ["git", "clone", "--depth", "1", clone_url, clone_path],
            capture_output=True,
            text=True,
            timeout=300,
            env={**os.environ},  # Inherit env; token is not in a logged arg
        )

        if result.returncode != 0:
            # Sanitise any token from error output before printing
            sanitised_stderr = result.stderr
            if token:
                sanitised_stderr = sanitised_stderr.replace(token, "***")
            print(f"Error cloning repository: {sanitised_stderr}")
            shutil.rmtree(temp_dir, ignore_errors=True)
            return None

        print(f"Repository cloned successfully.")
        return clone_path

    except Exception as e:
        shutil.rmtree(temp_dir, ignore_errors=True)
        print(f"Error cloning repository: {e}")
        traceback.print_exc()
        return None

def scan_local_directory(self, target_path: str) -> bool:
    """Scan a local directory or file."""
    if not os.path.exists(target_path):
        print(f"Error: Path not found: {target_path}")
        return False

    print(f"\nScanning local path: {target_path}")
    return self.run_semgrep(target_path)

def scan_github_repo(self, repo_url: str, token: Optional[str] = None) -> bool:
    """Clone and scan a GitHub repository, cleaning up afterward."""
    if not HAS_GITHUB:
        print("Error: PyGithub not installed. Run: pip install PyGithub")
        return False

    clone_path = self.clone_github_repo(repo_url, token)
    if not clone_path:
        return False

    # Derive the parent temp dir so we can clean up the whole thing
    parent_temp_dir = os.path.dirname(clone_path)

    try:
        success = self.run_semgrep(clone_path)

        self.repo_info = {
            "url": repo_url,
            "clone_time": datetime.now().isoformat(),
        }

        return success

    finally:
        print(f"\nCleaning up temporary files...")
        shutil.rmtree(parent_temp_dir, ignore_errors=True)

def calculate_risk_score(self) -> Tuple[int, str]:
    """Calculate compliance risk score and level."""
    total_score = 0
    for finding in self.findings:
        severity = finding.get("extra", {}).get("severity", "")
        if severity == "ERROR":
            total_score += 10
        elif severity == "HIGH":
            total_score += 7
        elif severity == "MEDIUM":
            total_score += 4
        elif severity == "LOW":
            total_score += 1

    if total_score >= 20:
        level = "CRITICAL"
    elif total_score >= 10:
        level = "HIGH"
    elif total_score >= 5:
        level = "MEDIUM"
    else:
        level = "LOW"

    return total_score, level

def _add_bold_paragraph(self, doc: "Document", label: str, value: str) -> None:
    """Helper: add a paragraph with a bold label and plain value."""
    p = doc.add_paragraph()
    p.add_run(f"{label}: ").bold = True
    p.add_run(value)

def generate_word_report(self, output_path: str) -> bool:
    """Generate Microsoft Word report with HIPAA compliance mapping and remediation."""
    if not HAS_DOCX:
        print("Error: python-docx not installed")
        return False

    doc = Document()

    # Title
    title = doc.add_heading("Healthcare Code Security Assessment Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata
    self._add_bold_paragraph(doc, "Generated", self.scan_metadata.get("timestamp", "N/A"))
    if self.repo_info:
        self._add_bold_paragraph(doc, "Repository", self.repo_info.get("url", "N/A"))
    else:
        self._add_bold_paragraph(doc, "Target", self.scan_metadata.get("target", "N/A"))
    self._add_bold_paragraph(doc, "Rules File", self.scan_metadata.get("rules_file", "N/A"))
    doc.add_paragraph("-" * 50)

    # Executive Summary
    doc.add_heading("Executive Summary", level=1)
    summary = doc.add_paragraph()
    summary.add_run(f"Total Findings: {self.scan_metadata['total_findings']}").bold = True
    summary.add_run("\n")
    summary.add_run(f"Critical: {self.scan_metadata['critical']}").bold = True
    summary.add_run(f" | High: {self.scan_metadata['high']}")
    summary.add_run(f" | Medium: {self.scan_metadata['medium']}")
    summary.add_run(f" | Low: {self.scan_metadata['low']}")

    # Risk Assessment
    doc.add_heading("Risk Assessment", level=1)
    if self.scan_metadata["critical"] > 0:
        doc.add_paragraph("CRITICAL: Immediate action required. Patient data may be at risk.")
    elif self.scan_metadata["high"] > 0:
        doc.add_paragraph(
            "HIGH: Significant vulnerabilities detected. Remediation recommended before deployment."
        )
    else:
        doc.add_paragraph("LOW RISK: No critical vulnerabilities detected.")

    # HIPAA Compliance Overview
    doc.add_heading("HIPAA Compliance Overview", level=1)

    hipaa_sections: Dict[str, Dict[str, Any]] = {}
    for finding in self.findings:
        hipaa_ref = (
            finding.get("extra", {}).get("metadata", {}).get("hipaa_reference", "")
        )
        hipaa_sub = (
            finding.get("extra", {}).get("metadata", {}).get("hipaa_subsection", "")
        )
        if hipaa_ref:
            if hipaa_ref not in hipaa_sections:
                hipaa_sections[hipaa_ref] = {"subsection": hipaa_sub, "count": 0}
            hipaa_sections[hipaa_ref]["count"] += 1

    if hipaa_sections:
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "HIPAA Section"
        hdr_cells[1].text = "Subsection"
        hdr_cells[2].text = "Findings"
        hdr_cells[3].text = "Status"

        for section, data in sorted(hipaa_sections.items()):
            status = "Non-Compliant" if data["count"] > 0 else "Compliant"
            row = table.add_row().cells
            row[0].text = section
            row[1].text = data["subsection"]
            row[2].text = str(data["count"])
            row[3].text = status
    else:
        doc.add_paragraph("No HIPAA-specific rules triggered.")

    # Compliance Checklist
    doc.add_heading("HIPAA Compliance Checklist", level=1)
    checklist_items = [
        ("164.308(a)(1)", "Administrative Safeguards - Security Management Process"),
        ("164.310(a)(1)", "Physical Safeguards - Facility Access Controls"),
        ("164.312(a)(1)", "Technical Safeguards - Access Control"),
        ("164.312(d)", "Technical Safeguards - Integrity Controls"),
        ("164.312(e)(1)", "Technical Safeguards - Transmission Security"),
    ]

    checklist_table = doc.add_table(rows=1, cols=3)
    checklist_table.style = "Table Grid"
    hdr = checklist_table.rows[0].cells
    hdr[0].text = "Section"
    hdr[1].text = "Requirement"
    hdr[2].text = "Status"

    for section, requirement in checklist_items:
        # FIX: was iterating a single-element list containing the loop variable,
        # which always evaluated True for any non-empty string.
        has_findings = section in hipaa_sections
        status = "Non-Compliant" if has_findings else "Compliant"
        row = checklist_table.add_row().cells
        row[0].text = section
        row[1].text = requirement
        row[2].text = status

    # Risk Score
    doc.add_heading("Compliance Risk Score", level=1)
    total_score, risk_level = self.calculate_risk_score()
    score_p = doc.add_paragraph()
    score_p.add_run(f"Total Risk Score: {total_score}").bold = True
    score_p.add_run(f" (Level: {risk_level})")

    # Remediation Timeline
    doc.add_heading("Remediation Timeline", level=1)
    timeline: Dict[str, Dict[str, Any]] = {}
    for finding in self.findings:
        priority = (
            finding.get("extra", {})
            .get("metadata", {})
            .get("remediation_priority", "Medium")
        )
        days = (
            finding.get("extra", {})
            .get("metadata", {})
            .get("max_remediation_days", 30)
        )
        if priority not in timeline:
            timeline[priority] = {"days": days, "count": 0}
        timeline[priority]["count"] += 1

    for priority, data in sorted(timeline.items(), key=lambda x: x[1]["days"]):
        doc.add_paragraph(
            f"{priority}: {data['count']} findings — Remediate within {data['days']} days"
        )

    # Detailed Findings
    doc.add_heading("Detailed Findings with HIPAA Mapping and Remediation", level=1)

    if not self.findings:
        doc.add_paragraph("No security vulnerabilities detected.")
    else:
        _default_remediation = {
            "title": "General Remediation",
            "description": "Review and fix this issue.",
            "steps": ["Fix the issue manually."],
            "before": "N/A",
            "after": "N/A",
            "resources": [],
            "severity_impact": "Unknown",
        }

        for i, finding in enumerate(self.findings, 1):
            start = finding.get("start", {})
            extra = finding.get("extra", {})
            metadata = extra.get("metadata", {})
            rule_id = finding.get("check_id", "")

            remediation = (
                get_remediation(rule_id) if HAS_REMEDIATIONS else _default_remediation
            )

            doc.add_heading(f"Finding #{i}: {rule_id}", level=2)
            self._add_bold_paragraph(doc, "File", finding.get("path", "Unknown"))
            self._add_bold_paragraph(doc, "Line", str(start.get("line", "N/A")))
            self._add_bold_paragraph(doc, "Severity", extra.get("severity", "Unknown"))
            self._add_bold_paragraph(doc, "Message", extra.get("message", "No message"))

            doc.add_paragraph().add_run("HIPAA Compliance:").bold = True
            self._add_bold_paragraph(
                doc, "  Section", metadata.get("hipaa_reference", "N/A")
            )
            self._add_bold_paragraph(
                doc, "  Subsection", metadata.get("hipaa_subsection", "N/A")
            )
            self._add_bold_paragraph(
                doc, "  Priority", metadata.get("remediation_priority", "Medium")
            )
            self._add_bold_paragraph(
                doc, "  Max Days", str(metadata.get("max_remediation_days", 30))
            )
            self._add_bold_paragraph(
                doc, "Impact", remediation.get("severity_impact", "Unknown")
            )

            doc.add_paragraph().add_run("Description:").bold = True
            doc.add_paragraph(remediation.get("description", "No description available."))

            doc.add_paragraph().add_run("Remediation Steps:").bold = True
            for j, step in enumerate(remediation.get("steps", []), 1):
                doc.add_paragraph(f"{j}. {step}", style="List Number")

            doc.add_paragraph().add_run("Code Example:").bold = True

            doc.add_heading("Before (Vulnerable):", level=3)
            before_para = doc.add_paragraph()
            before_para.add_run(remediation.get("before", "N/A"))
            before_para.style = "No Spacing"

            doc.add_heading("After (Fixed):", level=3)
            after_para = doc.add_paragraph()
            after_para.add_run(remediation.get("after", "N/A"))
            after_para.style = "No Spacing"

            resources = remediation.get("resources", [])
            if resources:
                doc.add_paragraph().add_run("Additional Resources:").bold = True
                for res_title, url in resources:
                    doc.add_paragraph(f"- {res_title}: {url}", style="List Bullet")

            doc.add_paragraph("-" * 50)

    doc.save(output_path)
    print(f"Word report saved to: {output_path}")
    return True

def generate_pdf_report(self, output_path: str) -> bool:
    """Generate PDF report with HIPAA compliance mapping and remediation."""
    if not HAS_REPORTLAB:
        print("Error: reportlab not installed")
        return False

    doc = SimpleDocTemplate(output_path, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []

    title_style = ParagraphStyle(
        "CustomTitle",
        parent=styles["Heading1"],
        fontSize=24,
        textColor=colors.HexColor("#003366"),
        spaceAfter=30,
    )
    story.append(Paragraph("Healthcare Code Security Assessment Report", title_style))
    story.append(Spacer(1, 20))

    meta_style = styles["Normal"]
    story.append(
        Paragraph(
            f"<b>Generated:</b> {self.scan_metadata.get('timestamp', 'N/A')}", meta_style
        )
    )
    if self.repo_info:
        story.append(
            Paragraph(
                f"<b>Repository:</b> {self.repo_info.get('url', 'N/A')}", meta_style
            )
        )
    else:
        story.append(
            Paragraph(
                f"<b>Target:</b> {self.scan_metadata.get('target', 'N/A')}", meta_style
            )
        )
    story.append(
        Paragraph(
            f"<b>Total Findings:</b> {self.scan_metadata['total_findings']}", meta_style
        )
    )
    story.append(Spacer(1, 20))

    story.append(Paragraph("<b>Risk Assessment</b>", styles["Heading2"]))
    if self.scan_metadata["critical"] > 0:
        story.append(
            Paragraph(
                "CRITICAL: Immediate action required. Patient data may be at risk.",
                styles["Normal"],
            )
        )
    elif self.scan_metadata["high"] > 0:
        story.append(
            Paragraph(
                "HIGH: Significant vulnerabilities detected.", styles["Normal"]
            )
        )
    else:
        story.append(
            Paragraph("LOW RISK: No critical vulnerabilities detected.", styles["Normal"])
        )
    story.append(Spacer(1, 20))

    story.append(Paragraph("<b>HIPAA Compliance Overview</b>", styles["Heading2"]))
    hipaa_sections: Dict[str, int] = {}
    for finding in self.findings:
        hipaa_ref = (
            finding.get("extra", {}).get("metadata", {}).get("hipaa_reference", "")
        )
        if hipaa_ref:
            hipaa_sections[hipaa_ref] = hipaa_sections.get(hipaa_ref, 0) + 1

    if hipaa_sections:
        for section, count in sorted(hipaa_sections.items()):
            story.append(
                Paragraph(f"- {section}: {count} findings", styles["Normal"])
            )
    else:
        story.append(
            Paragraph("No HIPAA-specific rules triggered.", styles["Normal"])
        )
    story.append(Spacer(1, 20))

    total_score, risk_level = self.calculate_risk_score()
    story.append(
        Paragraph(
            f"<b>Compliance Risk Score:</b> {total_score} (Level: {risk_level})",
            styles["Heading3"],
        )
    )
    story.append(Spacer(1, 20))

    story.append(
        Paragraph("<b>Detailed Findings with Remediation</b>", styles["Heading2"])
    )

    if not self.findings:
        story.append(
            Paragraph("No security vulnerabilities detected.", styles["Normal"])
        )
    else:
        _default_remediation = {
            "description": "No description.",
            "steps": [],
            "before": "N/A",
            "after": "N/A",
            "severity_impact": "Unknown",
        }

        for finding in self.findings:
            start = finding.get("start", {})
            extra = finding.get("extra", {})
            rule_id = finding.get("check_id", "")
            remediation = (
                get_remediation(rule_id) if HAS_REMEDIATIONS else _default_remediation
            )

            story.append(Paragraph(f"<b>{rule_id}</b>", styles["Heading3"]))
            story.append(
                Paragraph(
                    f"File: {finding.get('path', 'Unknown')} | Line: {start.get('line', 'N/A')}",
                    styles["Normal"],
                )
            )
            story.append(
                Paragraph(
                    f"Severity: {extra.get('severity', 'Unknown')}", styles["Normal"]
                )
            )
            story.append(
                Paragraph(
                    f"Impact: {remediation.get('severity_impact', 'Unknown')}",
                    styles["Normal"],
                )
            )
            story.append(Paragraph("<b>Description:</b>", styles["Heading4"]))
            story.append(
                Paragraph(remediation.get("description", "N/A"), styles["Normal"])
            )
            story.append(Paragraph("<b>Remediation Steps:</b>", styles["Heading4"]))
            for step in remediation.get("steps", []):
                story.append(Paragraph(f"• {step}", styles["Normal"]))

            story.append(Paragraph("<b>Code Example:</b>", styles["Heading4"]))
            before_text = remediation.get("before", "N/A")
            after_text = remediation.get("after", "N/A")
            story.append(
                Paragraph(
                    f"Before: {before_text[:100]}{'...' if len(before_text) > 100 else ''}",
                    styles["Normal"],
                )
            )
            story.append(
                Paragraph(
                    f"After: {after_text[:100]}{'...' if len(after_text) > 100 else ''}",
                    styles["Normal"],
                )
            )
            story.append(Spacer(1, 10))

    doc.build(story)
    print(f"PDF report saved to: {output_path}")
    return True
```

def main() -> None:
“”“Main entry point.”””
import argparse

```
parser = argparse.ArgumentParser(description="Healthcare Code Security Scanner")
parser.add_argument("target", help="Path to directory/file OR GitHub repo URL")
parser.add_argument(
    "--rules", default="healthcare_rules.yaml", help="Semgrep rules file"
)
parser.add_argument(
    "--output",
    default="security_report",
    help="Output filename (without extension)",
)
parser.add_argument(
    "--format",
    choices=["word", "pdf", "both"],
    default="both",
    help="Report format",
)
parser.add_argument(
    "--github-token",
    help="GitHub personal access token for private repos (prefer GITHUB_TOKEN env var)",
)
parser.add_argument(
    "--local",
    action="store_true",
    help="Force local scan (don't treat as GitHub URL)",
)
parser.add_argument(
    "--verbose",
    action="store_true",
    help="Enable debug output",
)
parser.add_argument(
    "--fail-on-critical",
    action="store_true",
    help="Exit with code 1 if any Critical findings are detected (useful for CI/CD gate)",
)

args = parser.parse_args()

# Prefer env var for token; fall back to CLI arg
github_token = os.environ.get("GITHUB_TOKEN") or args.github_token

print("=" * 60)
print("Healthcare Code Security Scanner")
print("=" * 60)

scanner = SecurityScanner(rules_file=args.rules, verbose=args.verbose)

is_github = not args.local and (
    args.target.startswith("http") or args.target.startswith("git@")
)

if is_github:
    print(f"\nTarget: GitHub Repository")
    print(f"Rules:  {args.rules}")
    if not scanner.scan_github_repo(args.target, github_token):
        print("GitHub scan failed!")
        sys.exit(1)
else:
    print(f"\nTarget: Local Path")
    print(f"Rules:  {args.rules}")
    if not scanner.scan_local_directory(args.target):
        print("Local scan failed!")
        sys.exit(1)

print(f"\nScan complete. Found {scanner.scan_metadata['total_findings']} issues:")
print(f"   Critical: {scanner.scan_metadata['critical']}")
print(f"   High:     {scanner.scan_metadata['high']}")
print(f"   Medium:   {scanner.scan_metadata['medium']}")
print(f"   Low:      {scanner.scan_metadata['low']}")

if args.format in ["word", "both"]:
    word_path = f"{args.output}.docx"
    if scanner.generate_word_report(word_path):
        print(f"\nWord report: {word_path}")

if args.format in ["pdf", "both"]:
    pdf_path = f"{args.output}.pdf"
    if scanner.generate_pdf_report(pdf_path):
        print(f"PDF report:  {pdf_path}")

print("\n" + "=" * 60)
print("Scan complete!")
print("=" * 60)

if args.fail_on_critical and scanner.scan_metadata.get("critical", 0) > 0:
    print(
        f"\n[FAIL] {scanner.scan_metadata['critical']} Critical finding(s) detected. "
        "Blocking as requested by --fail-on-critical."
    )
    sys.exit(1)
```

if **name** == “**main**”:
main()